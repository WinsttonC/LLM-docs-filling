from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat
# from langchain.llms.gigachat import GigaChat
from prompts import extraction_prompt
from utils import shorten_blanks
from docx import Document
from dotenv import load_dotenv
import os 
import warnings

load_dotenv()
warnings.filterwarnings("ignore")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 4

chat = GigaChat(credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False) #model='GigaChat-Pro'
doc = Document('dogovor_arenda.docx')

final_text = []
paragraphs = [shorten_blanks(p.text) for p in doc.paragraphs]

length = len(paragraphs)
print('======== НАЧАЛО ОБРАБОТКИ ДОКУМЕНТА ========')
print(f'Нужно обработать {length} строк')
for i in range(0, length, BATCH_SIZE):
    batch = paragraphs[i:min(i + BATCH_SIZE, length)]
    input_text = '\n'.join(batch)
    input_text = f'Строка с пропуском: {input_text}\nОтвет:'
    messages = [SystemMessage(content=extraction_prompt)]
    messages.append(HumanMessage(content=input_text))
    res = chat(messages)
    final_text.append(res.content)
    # messages.append(res) # Сохранение истории запросов
    print(f"\rОбработано {i} из {length}")



# for i in range(len(final_text)):
#     batch = final_text[i]
#     if '_' in batch:   
#         messages = [SystemMessage(content=extraction_prompt)]
#         messages.append(HumanMessage(content=batch))
#         res = chat(messages)
#         final_text[i] = res.content
#         # messages.append(res) # Сохранение истории запросов
#         print(f'Запрос \n{batch}')
#         print(f'Ответ\n\n{res.content}')

new_doc = Document()

# Добавляем абзац с текстом в документ\
for text_block in final_text:
    new_doc.add_paragraph(text_block)

# Сохраняем документ
new_doc.save('processed_arenda.docx')

print("ПОИСК СУЩНОСТЕЙ В ДОКУМЕНТЕ УСПЕШНО ЗАВЕРШЕН.")