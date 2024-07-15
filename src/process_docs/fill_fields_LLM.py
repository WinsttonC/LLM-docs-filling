from prompts import filling_prompt
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat
from docx import Document
import json
from dotenv import load_dotenv
import os
import re
import warnings
load_dotenv()

warnings.filterwarnings("ignore")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 4

chat = GigaChat(credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False) #model='GigaChat-Pro'

document = Document('processed_arenda.docx')
paragraphs = [p.text for p in document.paragraphs]

with open("schema.json", "r", encoding='utf-8') as f:
    schema = json.load(f)

filled_text = []

print('======== НАЧАЛО ОБРАБОТКИ ДОКУМЕНТА ========')
print(f'Нужно заполнить {len(schema)} пропусков.')
length = len(paragraphs)
for i in range(0, length, BATCH_SIZE):
    batch = '\n'.join(paragraphs[i:min(i + BATCH_SIZE, length)])

    user_data = []

    pattern = r'\[\[(.*?)\]\]'
    matches = re.findall(pattern, batch)
    for elem in matches:
        
        data = schema[elem].get('user_answer', 'Данные неизвестны')
        hint = f'{elem}: {data}'
        user_data.append(hint)

    input_text = f'Строка документа:\n{batch}\n\n Данные пользователя:{user_data}\n\n Ответ:'
    messages = [SystemMessage(content=filling_prompt)]
    messages.append(HumanMessage(content=input_text))
    res = chat(messages)
    filled_text.append(res.content)

    print(f"\Заполнено {i} из {len(schema)} пропусков")
    # # messages.append(res) # Сохранение истории запросов
    # print(f'Запрос \n{input_text}')
    # print(f'Ответ\n\n{res.content}')

final_document = Document()

# Добавляем абзац с текстом в документ\
for text_block in filled_text:
    # print(text_block)
    final_document.add_paragraph(text_block)

# Сохраняем документ
final_document.save('final_arenda.docx')

print("ДОКУМЕНТ ЗАПОЛНЕН.")