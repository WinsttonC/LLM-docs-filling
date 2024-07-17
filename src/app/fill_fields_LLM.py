from prompts import filling_prompt
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat
from docx import Document
import json
from dotenv import load_dotenv
import os
import re
import warnings
load_dotenv()

warnings.filterwarnings("ignore")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 1

chat = GigaChat(credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False) #model='GigaChat-Pro'
doc_path = os.getenv("DOCUMENTS_PATH")




def fill_fields_LLM(doc_name: str, fields_dict: dict):
    path_to_docx = f'{doc_path}/documents/{doc_name}.docx'

    document = Document(path_to_docx)
    paragraphs = [p.text for p in document.paragraphs]
    final_document = Document()
    # filled_text = []
    length = len(paragraphs)
    for i in range(0, length, BATCH_SIZE):
        batch = '\n'.join(paragraphs[i:min(i + BATCH_SIZE, length)])

        user_data = []

        pattern = r'\[\[(.*?)\]\]'
        matches = re.findall(pattern, batch)
        if matches:
            for elem in matches:
                
                data = fields_dict.get(elem, 'Данные неизвестны')
                hint = f'{elem}: {data}'
                user_data.append(hint)
            user_data = '\n'.join(user_data)
            input_text = f'Строка документа:\n{batch}\n\n Данные пользователя:{user_data}\n\n Ответ:'
            messages = [SystemMessage(content=filling_prompt)]
            messages.append(HumanMessage(content=input_text))
            res = chat(messages)

            final_document.add_paragraph(res.content)
            # filled_text.append(res.content)
        else:
            final_document.add_paragraph(batch)
            # filled_text.append(batch)
        

        # for text_block in filled_text:
        #     final_document.add_paragraph(text_block)

        final_document.save(f'{doc_path}/user_docs/{doc_name}.docx')
