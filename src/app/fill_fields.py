import os
import re
import warnings

from docx import Document
from dotenv import load_dotenv
from utils import replace_substrings
from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat

from prompts import filling_prompt

load_dotenv()

warnings.filterwarnings("ignore")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 1

chat = GigaChat(
    credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False
)  # model='GigaChat-Pro'
doc_path = os.getenv("DOCUMENTS_PATH")



def fill_fields_rules(doc_name: str, fields_dict: dict):
    path_to_docx = f"{doc_path}/documents/{doc_name}.docx"

    str_document = Document(path_to_docx)
    str_document = [p.text for p in str_document.paragraphs]
    for i in range(len(str_document)):
        str_document[i] = replace_substrings(str_document[i], fields_dict)

    final_document = Document()

    for text_block in str_document:
        final_document.add_paragraph(text_block)

    final_document.save(f"{doc_path}/user_docs/{doc_name}.docx")





def fill_fields_LLM(doc_name: str, fields_dict: dict):
    path_to_docx = f"{doc_path}/documents/{doc_name}.docx"

    document = Document(path_to_docx)
    paragraphs = [p.text for p in document.paragraphs]
    final_document = Document()

    length = len(paragraphs)
    for i in range(0, length, BATCH_SIZE):
        batch = "\n".join(paragraphs[i : min(i + BATCH_SIZE, length)])

        user_data = []

        pattern = r"\[\[(.*?)\]\]"
        matches = re.findall(pattern, batch)
        if matches:
            for elem in matches:
                data = fields_dict.get(elem, "Данные неизвестны")
                hint = f"{elem}: {data}"
                user_data.append(hint)
            user_data = "\n".join(user_data)
            input_text = f"Строка документа:\n{batch}\n\n Данные пользователя:{user_data}\n\n Ответ:"
            messages = [SystemMessage(content=filling_prompt)]
            messages.append(HumanMessage(content=input_text))
            res = chat(messages)

            final_document.add_paragraph(res.content)
        else:
            final_document.add_paragraph(batch)

        final_document.save(f"{doc_path}/user_docs/{doc_name}.docx")