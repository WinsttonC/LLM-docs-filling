from langchain.schema import HumanMessage, SystemMessage
from langchain_community.chat_models.gigachat import GigaChat
# from langchain.llms.gigachat import GigaChat
from prompts import extraction_prompt
from utils import shorten_blanks
from docx import Document
from dotenv import load_dotenv
import os 
import re
import warnings
load_dotenv()
warnings.filterwarnings("ignore")

doc_path = os.getenv("DOCUMENTS_PATH")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 2
chat = GigaChat(credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False) #model='GigaChat-Pro'

def process_doc(doc_name, new_doc=False):
    if new_doc:
        file_path = f'{doc_path}/raw_docs/{doc_name}.docx'
    else:
        file_path = f'{doc_path}/documents/{doc_name}.docx'

    doc = Document(file_path)

    final_text = []
    paragraphs = [shorten_blanks(p.text) for p in doc.paragraphs]
    pattern = r'_{2,}'
    
    length = len(paragraphs)
    for i in range(0, length, BATCH_SIZE):
        batch = paragraphs[i:min(i + BATCH_SIZE, length)]
        input_text = '\n'.join(batch)
        match = re.search(pattern, input_text)
        if match:
            input_text = f'Строка:\n {input_text}\nОтвет:'
            messages = [SystemMessage(content=extraction_prompt)]
            messages.append(HumanMessage(content=input_text))
            res = chat(messages)
            final_text.append(res.content)
            print('В цикле с LLM:')
            print('Строка: \n', input_text)
            print('Результат: \n', res.content)
            print('========================================\n')
        else:
            print('В цикле REGEX:')
            # print('Строка: \n', input_text)
            print('Результат: \n', input_text)
            print('========================================\n')
            final_text.append(input_text)

    save_path = f'{doc_path}/documents/{doc_name}.docx'
    new_doc = Document()

    for text_block in final_text:
        new_doc.add_paragraph(text_block)

    new_doc.save(save_path)
    
