import os
import re

from .agents.find_fields_agent import find_fields
from docx import Document
from dotenv import load_dotenv

load_dotenv()


doc_path = os.getenv("DOCUMENTS_PATH")
BATCH_SIZE = 1


def is_valid_string(input_string):
    '''
    Функция для проверки, является ли строка пропуском(__).

    Parameters
    ----------
    input_string : str

    Returns
    ---------
    True/False
    '''

    pattern = r"^[_. ]+$"
    return not bool(re.match(pattern, input_string))


def shorten_blanks(input_string):
    '''
    Заменяет длинные пропуски(__) в документе на пропуски
    фиксированной длины, чтобы упростить поиск в дальнейшем 
    и избавиться от строк, состоящий только из __.

    Parameters
    ----------
    input_string : str
        Любая строка из документа.
    Returns
    ---------
    output_string : str
        Обновленная строка с пропусками или пустая строка.
    '''

    if is_valid_string(input_string):
        pattern = r"(_{2,})"
        output_string = re.sub(pattern, "____________", input_string)

        return output_string
    else:
        return ""


def create_fields_template(doc_name, new_doc=False):
    '''
    Ищет пропуски в неразмеченном документе и заменяет их на шаблон 
    с описанием, например [[ФИО истца]].
        1. Разбивает текст на блоки по абзацам
        2. Каждый абзац с пропуском передается в LLM
        3. Блоки собираются в новый .docx документ

    Parameters
    ----------
    doc_name : str
        Название документа.
    new_doc: bool
        True - документ уже есть в базе данных.
        False - новый документ, загруженный пользователем.
    Returns
    ---------
    None
        Сохраняет документ с замененными пропусками
        в хранилище с размеченными документами.
    '''
    
    if new_doc:
        file_path = f"{doc_path}/raw_docs/{doc_name}.docx"
    else:
        file_path = f"{doc_path}/documents/{doc_name}.docx"

    doc = Document(file_path)

    final_text = []
    paragraphs = [shorten_blanks(p.text) for p in doc.paragraphs]

    length = len(paragraphs)
    for i in range(0, length, BATCH_SIZE):
        batch = paragraphs[i : min(i + BATCH_SIZE, length)]
        input_text = "\n".join(batch)
        
        if '__' in input_text:
            
            input_text = f"Строка:\n {input_text}\nОтвет:"
            content = find_fields(input_text)
            if '__' in content:
                content = find_fields(input_text)

            final_text.append(content)
        
        else:
            final_text.append(input_text)

    save_path = f"{doc_path}/documents/{doc_name}.docx"
    new_doc = Document()

    for text_block in final_text:
        new_doc.add_paragraph(text_block)

    new_doc.save(save_path)
