import os
import re
import warnings

from .agents.filling_fields_agent import fill_fields
from docx import Document
from dotenv import load_dotenv
from langchain_community.chat_models.gigachat import GigaChat

load_dotenv()

warnings.filterwarnings("ignore")
GIGACHAT_CLIENT_SECRET = os.getenv("GIGACHAT_CLIENT_SECRET_B64")
BATCH_SIZE = 1

chat = GigaChat(
    credentials=GIGACHAT_CLIENT_SECRET, verify_ssl_certs=False
) 
doc_path = os.getenv("DOCUMENTS_PATH")


def replace_substrings(input_string, fields_dict):
    '''
    Заменяет подстроки вида [[Описание пропуска]] на
    данные из словаря с ответами пользователя.

    Parameters
    ----------
    input_string : str
        Строка с размеченными пропусками.
    fields_dict: dict
        Словарь с данными пользователя в формате:
        {'ФИО заявителя': 'Иванов Иван Иванович'}
    Returns
    ---------
    result_string : str
        Строка, заполненная данными пользователя.
    '''

    pattern = r"\[\[(.*?)\]\]"

    # Функция замены подстрок
    def replace_value(match):
        key = match.group(1)
        return fields_dict.get(key, f"[[{key}]]")

    result_string = re.sub(pattern, replace_value, input_string)

    return result_string


def fill_fields_rules(doc_name: str, fields_dict: dict):
    '''
    Заменяет все подстроки вида [[Описание пропуска]] в документе 
    с использованием регулярных выражений. 

    Parameters
    ----------
    doc_name : str
        Документ с размеченными пропусками.
    fields_dict: dict
        Словарь с данными пользователя в формате:
        {'ФИО заявителя': 'Иванов Иван Иванович'}
    Returns
    ---------
    None
        Сохраняет заполненный документ с хранилище с данными пользователя.
    '''

    path_to_docx = f"{doc_path}/documents/{doc_name}.docx"

    str_document = Document(path_to_docx)
    str_document = [p.text for p in str_document.paragraphs]
    for i in range(len(str_document)):
        str_document[i] = replace_substrings(str_document[i], fields_dict)

    final_document = Document()

    for text_block in str_document:
        final_document.add_paragraph(text_block)

    final_document.save(f"{doc_path}/user_docs/{doc_name}.docx")


def fill_fields_LLM(doc_name: str, fields_dict: dict):
    '''
    Заменяет все подстроки вида [[Описание пропуска]]]
    в документе с использованием LLM. 
        1. Разбивает документ на блоки по абзацам
        2. Заполняет пропуски в каждом абзаце,
        используя данные пользователя
        2*. Если пропуски остались, повторяет 2 шаг
        3. Сохраняет заполненный документ в хранилище 
        с файлами пользователя.

    Parameters
    ----------
    doc_name : str
        Документ с размеченными пропусками.
    fields_dict: dict
        Словарь с данными пользователя в формате:
        {'ФИО заявителя': 'Иванов Иван Иванович'}
    Returns
    ---------
    None
        Сохраняет заполненный документ с хранилище с файлами пользователя.
    '''

    path_to_docx = f"{doc_path}/documents/{doc_name}.docx"

    document = Document(path_to_docx)
    paragraphs = [p.text for p in document.paragraphs]
    final_document = Document()

    length = len(paragraphs)
    for i in range(0, length, BATCH_SIZE):
        batch = "\n".join(paragraphs[i : min(i + BATCH_SIZE, length)])

        user_data = []

        pattern = r"\[\[(.*?)\]\]"
        matches = re.findall(pattern, batch)
        if matches:
            for elem in matches:
                data = fields_dict.get(elem, "Данные неизвестны")
                hint = f"{elem}: {data}"
                user_data.append(hint)
            user_data = "\n".join(user_data)
            input_text = f"Строка документа:\n{batch}\n\n Данные пользователя:{user_data}\n\n Ответ:"
            content = fill_fields(input_text)

            final_document.add_paragraph(content)
        else:
            final_document.add_paragraph(batch)

        final_document.save(f"{doc_path}/user_docs/{doc_name}.docx")
