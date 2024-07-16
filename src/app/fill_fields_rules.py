from utils import replace_substrings
from docx import Document
from dotenv import load_dotenv
import os
load_dotenv()

doc_path = os.getenv("DOCUMENTS_PATH")

def fill_fields_rules(doc_name: str, fields_dict: dict):
    path_to_docx = f'{doc_path}/documents/{doc_name}.docx'

    str_document = Document(path_to_docx)
    str_document = [p.text for p in str_document.paragraphs]
    for i in range(len(str_document)):
        str_document[i] = replace_substrings(str_document[i], fields_dict)
        
    final_document = Document()

    # Добавляем абзац с текстом в документ\
    for text_block in str_document:
        final_document.add_paragraph(text_block)

    # Сохраняем документ
    final_document.save(f'{doc_path}/user_docs/{doc_name}.docx')